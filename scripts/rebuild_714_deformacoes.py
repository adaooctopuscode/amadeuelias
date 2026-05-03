#!/usr/bin/env python3
"""
Substitui apenas os blocos de figuras da subseção DEFORMAÇÕES (vigas protendidas),
mantendo o título e os parágrafos introdutórios até «Combinação 2: 1,0* Cargas acidentais».

Para cada tabuleiro e cada combinação (1 e 2):
  [imagem 5,5 in] [Figura N: ...] [Fonte: Software SCIA.] [imagem sem legenda]

Prefixos conforme pasta do WhatsApp (sem arquivo 00000260).
"""

from __future__ import annotations

import os
import re
from dataclasses import dataclass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from lxml import etree

ROOT = "/Users/carlos_adao/Documents/2AEngenharia/Amadeu_Elias"
DOC_PATH = os.path.join(ROOT, "MC-AMADEU ELIAS R0.docx")
IMG_DIR = os.path.join(ROOT, "WhatsApp Chat - PONTE AMADEU")

H_DEF = "DEFORMAÇÕES NAS VIGAS PROTENIDAS"
LAST_INTRO = "Combinação 2: 1,0* Cargas acidentais"
H_SUPER = "SUPERESTRUTURA"

# por tabuleiro: comb1 principal, comb1 extra, comb2 principal, comb2 extra
PREFIX_GROUPS = (
    ("00000248", "00000249", "00000250", "00000251"),
    ("00000252", "00000253", "00000254", "00000255"),
    ("00000256", "00000257", "00000258", "00000259"),
    ("00000261", "00000262", "00000263", "00000264"),
    ("00000265", "00000266", "00000267", "00000268"),
    ("00000269", "00000270", "00000271", "00000272"),
    ("00000273", "00000274", "00000275", "00000276"),
)


@dataclass
class ParaSnap:
    style_name: str | None
    alignment: object
    space_before: object
    space_after: object
    line_spacing: object
    line_spacing_rule: object
    left_indent: object
    right_indent: object
    first_line_indent: object
    keep_together: object
    keep_with_next: object
    run_bold: object
    run_italic: object
    run_font_name: object
    run_font_size: object


def _snap(p: Paragraph) -> ParaSnap:
    pf = p.paragraph_format
    r0 = p.runs[0] if p.runs else None
    return ParaSnap(
        style_name=p.style.name if p.style else None,
        alignment=pf.alignment,
        space_before=pf.space_before,
        space_after=pf.space_after,
        line_spacing=pf.line_spacing,
        line_spacing_rule=pf.line_spacing_rule,
        left_indent=pf.left_indent,
        right_indent=pf.right_indent,
        first_line_indent=pf.first_line_indent,
        keep_together=pf.keep_together,
        keep_with_next=pf.keep_with_next,
        run_bold=r0.font.bold if r0 else None,
        run_italic=r0.font.italic if r0 else None,
        run_font_name=r0.font.name if r0 else None,
        run_font_size=r0.font.size if r0 else None,
    )


def _apply_snap(p: Paragraph, s: ParaSnap):
    if s.style_name:
        p.style = s.style_name
    pf = p.paragraph_format
    pf.alignment = s.alignment
    pf.space_before = s.space_before
    pf.space_after = s.space_after
    pf.line_spacing = s.line_spacing
    pf.line_spacing_rule = s.line_spacing_rule
    pf.left_indent = s.left_indent
    pf.right_indent = s.right_indent
    pf.first_line_indent = s.first_line_indent
    pf.keep_together = s.keep_together
    pf.keep_with_next = s.keep_with_next


def _apply_run_snap(run, s: ParaSnap):
    run.bold = s.run_bold
    run.italic = s.run_italic
    run.font.name = s.run_font_name
    run.font.size = s.run_font_size


def _strip_paragraph_children(paragraph: Paragraph):
    for child in list(paragraph._p):
        if etree.QName(child).localname != "pPr":
            paragraph._p.remove(child)


def _find_image_file(prefix: str) -> str:
    for name in os.listdir(IMG_DIR):
        if name.startswith(prefix + "-") and name.lower().endswith((".jpg", ".jpeg", ".png")):
            return os.path.join(IMG_DIR, name)
    raise FileNotFoundError(f"Imagem com prefixo {prefix} não encontrada em {IMG_DIR}")


def _insert_paragraph_after(anchor: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    return Paragraph(new_p, anchor._parent)


def _renumber_figures(doc: Document, cap_snap: ParaSnap):
    fig_n = 1
    fig_re = re.compile(r"^Figura\s+(\d+)\s*:\s*(.*)$", re.DOTALL)

    for para in doc.paragraphs:
        raw = para.text.strip()
        if not raw.startswith("Figura"):
            continue
        m = fig_re.match(raw)
        if not m:
            continue
        suffix = m.group(2).strip()
        _strip_paragraph_children(para)
        _apply_snap(para, cap_snap)
        rr = para.add_run(f"Figura {fig_n}: {suffix}")
        _apply_run_snap(rr, cap_snap)
        fig_n += 1

    print(f"Legendas 'Figura' renumeradas até {fig_n - 1}.")


def rebuild_section():
    doc = Document(DOC_PATH)

    def_idx = intro_idx = super_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == H_DEF:
            def_idx = i
    if def_idx is None:
        raise RuntimeError(f"Título {H_DEF!r} não encontrado.")

    for i in range(def_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip() == LAST_INTRO:
            intro_idx = i
            break
    if intro_idx is None:
        raise RuntimeError(f"Parágrafo introdutório {LAST_INTRO!r} não encontrado.")

    # Há outra ocorrência de SUPERESTRUTURA antes desta subseção no documento;
    # usar apenas o primeiro Heading após o último parágrafo introdutório.
    for i in range(intro_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip() == H_SUPER:
            super_idx = i
            break
    if super_idx is None or super_idx <= intro_idx:
        raise RuntimeError(f"Título {H_SUPER!r} não encontrado após {LAST_INTRO!r}.")

    # Modelos a partir do primeiro bloco existente (antes de apagar)
    first_fig = None
    first_img = None
    for i in range(intro_idx + 1, super_idx):
        p = doc.paragraphs[i]
        if first_fig is None and p.text.strip().startswith("Figura "):
            first_fig = i
            break
    if first_fig is None:
        raise RuntimeError("Não foi possível localizar uma legenda modelo.")

    # Parágrafo de imagem: primeiro com blip antes da primeira legenda; senão, logo após Fonte.
    for i in range(intro_idx + 1, first_fig):
        if "blip" in doc.paragraphs[i]._p.xml:
            first_img = i
            break
    if first_img is None:
        for i in range(first_fig + 2, min(first_fig + 8, super_idx)):
            if "blip" in doc.paragraphs[i]._p.xml:
                first_img = i
                break
    if first_img is None:
        raise RuntimeError("Não foi possível localizar parágrafo de imagem modelo.")

    snap_img = _snap(doc.paragraphs[first_img])
    snap_cap = _snap(doc.paragraphs[first_fig])
    snap_fonte = _snap(doc.paragraphs[first_fig + 1])

    if snap_cap.alignment is None:
        snap_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if snap_fonte.alignment is None:
        snap_fonte.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if snap_img.alignment is None:
        snap_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

    to_remove = [doc.paragraphs[i]._element for i in range(intro_idx + 1, super_idx)]
    for el in to_remove:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)

    anchor = doc.paragraphs[intro_idx]

    def add_image_only(prefix: str):
        nonlocal anchor
        img_path = _find_image_file(prefix)
        p_img = _insert_paragraph_after(anchor)
        anchor = p_img
        _apply_snap(p_img, snap_img)
        _strip_paragraph_children(p_img)
        r = p_img.add_run()
        r.add_picture(img_path, width=Inches(5.5))

    def add_labeled_pair(prefix_img: str, caption_text: str):
        nonlocal anchor
        img_path = _find_image_file(prefix_img)
        p_img = _insert_paragraph_after(anchor)
        anchor = p_img
        _apply_snap(p_img, snap_img)
        _strip_paragraph_children(p_img)
        r = p_img.add_run()
        r.add_picture(img_path, width=Inches(5.5))

        p_cap = _insert_paragraph_after(anchor)
        anchor = p_cap
        _apply_snap(p_cap, snap_cap)
        _strip_paragraph_children(p_cap)
        rr = p_cap.add_run(f"Figura 0: {caption_text}")
        _apply_run_snap(rr, snap_cap)

        p_fonte = _insert_paragraph_after(anchor)
        anchor = p_fonte
        _apply_snap(p_fonte, snap_fonte)
        _strip_paragraph_children(p_fonte)
        rr2 = p_fonte.add_run("Fonte: Software SCIA.")
        _apply_run_snap(rr2, snap_fonte)

    for tab in range(1, 8):
        p1, x1, p2, x2 = PREFIX_GROUPS[tab - 1]
        add_labeled_pair(
            p1,
            f"Deformações nas vigas protendidas do Tabuleiro {tab} para a Combinação 1.",
        )
        add_image_only(x1)
        add_labeled_pair(
            p2,
            f"Deformações nas vigas protendidas do Tabuleiro {tab} para a Combinação 2.",
        )
        add_image_only(x2)

    _renumber_figures(doc, snap_cap)
    doc.save(DOC_PATH)


if __name__ == "__main__":
    rebuild_section()
