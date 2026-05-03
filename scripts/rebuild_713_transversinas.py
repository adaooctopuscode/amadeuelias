#!/usr/bin/env python3
"""
Reconstrói apenas a seção entre:
  - ESFORÇOS CARACTERÍSTICOS NAS TRANSVERSINAS
  - DEFORMAÇÕES NAS VIGAS PROTENIDAS
com 28 figuras (plano WhatsApp: 00000217–00000230 e 00000232–00000245).
Não altera texto fora desse intervalo, exceto renumeração global de legendas "Figura N:".
"""

from __future__ import annotations

import os
import re
from copy import deepcopy
from dataclasses import dataclass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from lxml import etree

WNS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

ROOT = "/Users/carlos_adao/Documents/2AEngenharia/Amadeu_Elias"
DOC_PATH = os.path.join(ROOT, "MC-AMADEU ELIAS R0.docx")
IMG_DIR = os.path.join(ROOT, "WhatsApp Chat - PONTE AMADEU")

# permanente M, permanente V, acidental M, acidental V
PREFIX_BY_TAB = {
    1: ("00000217", "00000224", "00000232", "00000239"),
    2: ("00000218", "00000225", "00000233", "00000240"),
    3: ("00000219", "00000226", "00000234", "00000241"),
    4: ("00000220", "00000227", "00000235", "00000242"),
    5: ("00000221", "00000228", "00000236", "00000243"),
    6: ("00000222", "00000229", "00000237", "00000244"),
    7: ("00000223", "00000230", "00000238", "00000245"),
}

H_START = "ESFORÇOS CARACTERÍSTICOS NAS TRANSVERSINAS"
# Título exatamente como no documento (ortografia existente)
H_END = "DEFORMAÇÕES NAS VIGAS PROTENIDAS"


@dataclass
class ParaSnap:
    style_name: str | None
    alignment: object
    space_before: object
    space_after: object
    line_spacing: object
    line_spacing_rule: object
    left_indent: object
    right_indent: object
    first_line_indent: object
    keep_together: object
    keep_with_next: object
    run_bold: object
    run_italic: object
    run_font_name: object
    run_font_size: object


def _snap(p: Paragraph) -> ParaSnap:
    pf = p.paragraph_format
    r0 = p.runs[0] if p.runs else None
    return ParaSnap(
        style_name=p.style.name if p.style else None,
        alignment=pf.alignment,
        space_before=pf.space_before,
        space_after=pf.space_after,
        line_spacing=pf.line_spacing,
        line_spacing_rule=pf.line_spacing_rule,
        left_indent=pf.left_indent,
        right_indent=pf.right_indent,
        first_line_indent=pf.first_line_indent,
        keep_together=pf.keep_together,
        keep_with_next=pf.keep_with_next,
        run_bold=r0.font.bold if r0 else None,
        run_italic=r0.font.italic if r0 else None,
        run_font_name=r0.font.name if r0 else None,
        run_font_size=r0.font.size if r0 else None,
    )


def _apply_snap(p: Paragraph, s: ParaSnap):
    if s.style_name:
        p.style = s.style_name
    pf = p.paragraph_format
    pf.alignment = s.alignment
    pf.space_before = s.space_before
    pf.space_after = s.space_after
    pf.line_spacing = s.line_spacing
    pf.line_spacing_rule = s.line_spacing_rule
    pf.left_indent = s.left_indent
    pf.right_indent = s.right_indent
    pf.first_line_indent = s.first_line_indent
    pf.keep_together = s.keep_together
    pf.keep_with_next = s.keep_with_next


def _apply_run_snap(run, s: ParaSnap):
    run.bold = s.run_bold
    run.italic = s.run_italic
    run.font.name = s.run_font_name
    run.font.size = s.run_font_size


def _ensure_p_pr(p_el):
    pPr = p_el.find(WNS + "pPr")
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_el.insert(0, pPr)
    return pPr


def _set_list_num_pr(paragraph: Paragraph, num_pr_template: etree._Element | None):
    if num_pr_template is None:
        return
    p_el = paragraph._p
    pPr = _ensure_p_pr(p_el)
    old = pPr.find(WNS + "numPr")
    if old is not None:
        pPr.remove(old)
    pPr.append(deepcopy(num_pr_template))


def _strip_paragraph_children(paragraph: Paragraph):
    for child in list(paragraph._p):
        if etree.QName(child).localname != "pPr":
            paragraph._p.remove(child)


def _find_image_file(prefix: str) -> str:
    for name in os.listdir(IMG_DIR):
        if name.startswith(prefix + "-") and name.lower().endswith((".jpg", ".jpeg", ".png")):
            return os.path.join(IMG_DIR, name)
    raise FileNotFoundError(f"Imagem com prefixo {prefix} não encontrada em {IMG_DIR}")


def _extract_num_pr(doc: Document, para_idx: int):
    p = doc.paragraphs[para_idx]._p
    pPr = p.find(WNS + "pPr")
    if pPr is None:
        return None
    np = pPr.find(WNS + "numPr")
    return deepcopy(np) if np is not None else None


def _insert_paragraph_after(anchor: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    return Paragraph(new_p, anchor._parent)


def _renumber_figures(doc: Document, cap_snap: ParaSnap):
    fig_n = 1
    fig_re = re.compile(r"^Figura\s+(\d+)\s*:\s*(.*)$", re.DOTALL)

    for para in doc.paragraphs:
        raw = para.text.strip()
        if not raw.startswith("Figura"):
            continue
        m = fig_re.match(raw)
        if not m:
            continue
        suffix = m.group(2).strip()
        _strip_paragraph_children(para)
        _apply_snap(para, cap_snap)
        rr = para.add_run(f"Figura {fig_n}: {suffix}")
        _apply_run_snap(rr, cap_snap)
        fig_n += 1

    print(f"Legendas 'Figura' renumeradas até {fig_n - 1}.")


def rebuild_section():
    doc = Document(DOC_PATH)

    sec_idx = None
    end_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == H_START:
            sec_idx = i
        elif sec_idx is not None and t == H_END:
            end_idx = i
            break

    if sec_idx is None or end_idx is None:
        raise RuntimeError(f"Títulos não encontrados: {H_START!r} / {H_END!r}")

    # Modelos no estado atual (antes de apagar o intervalo)
    num_pr_intro = _extract_num_pr(doc, 992)
    num_pr_sub = _extract_num_pr(doc, 995)
    snap_empty = _snap(doc.paragraphs[993])
    snap_img = _snap(doc.paragraphs[996])
    snap_cap = _snap(doc.paragraphs[997])
    snap_fonte = _snap(doc.paragraphs[998])

    if snap_cap.alignment is None:
        snap_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if snap_fonte.alignment is None:
        snap_fonte.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if snap_img.alignment is None:
        snap_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

    to_remove = [doc.paragraphs[i]._element for i in range(sec_idx + 1, end_idx)]
    for el in to_remove:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)

    anchor = doc.paragraphs[sec_idx]

    def add_blank():
        nonlocal anchor
        p = _insert_paragraph_after(anchor)
        anchor = p
        _apply_snap(p, snap_empty)
        _strip_paragraph_children(p)

    def add_list_line(text: str, num_pr):
        nonlocal anchor
        p = _insert_paragraph_after(anchor)
        anchor = p
        p.style = "List Paragraph"
        _strip_paragraph_children(p)
        p.add_run(text)
        _set_list_num_pr(p, num_pr)

    def add_figure_block(rel_img_prefix: str, caption_body: str):
        nonlocal anchor
        img_path = _find_image_file(rel_img_prefix)

        p_img = _insert_paragraph_after(anchor)
        anchor = p_img
        _apply_snap(p_img, snap_img)
        _strip_paragraph_children(p_img)
        r = p_img.add_run()
        r.add_picture(img_path, width=Inches(5.5))

        p_cap = _insert_paragraph_after(anchor)
        anchor = p_cap
        _apply_snap(p_cap, snap_cap)
        _strip_paragraph_children(p_cap)
        rr = p_cap.add_run(f"Figura 0: {caption_body}")
        _apply_run_snap(rr, snap_cap)

        p_fonte = _insert_paragraph_after(anchor)
        anchor = p_fonte
        _apply_snap(p_fonte, snap_fonte)
        _strip_paragraph_children(p_fonte)
        rr2 = p_fonte.add_run("Fonte: Software SCIA.")
        _apply_run_snap(rr2, snap_fonte)

    for tab in range(1, 8):
        pm, pc, am, ac = PREFIX_BY_TAB[tab]
        add_list_line(
            f"Diagramas de momentos e cortantes característicos nas transversinas do Tabuleiro {tab}:",
            num_pr_intro,
        )
        add_blank()
        add_list_line("Para as cargas permanentes", num_pr_sub)
        add_blank()
        add_figure_block(
            pm,
            f"Diagramas de momentos fletores nas transversinas do Tabuleiro {tab}.",
        )
        add_figure_block(
            pc,
            f"Diagramas de esforços cortantes nas transversinas do Tabuleiro {tab}.",
        )
        add_list_line("Para as cargas acidentais", num_pr_sub)
        add_blank()
        add_figure_block(
            am,
            f"Diagramas de momentos fletores nas transversinas do Tabuleiro {tab}.",
        )
        add_figure_block(
            ac,
            f"Diagramas de esforços cortantes nas transversinas do Tabuleiro {tab}.",
        )

    _renumber_figures(doc, snap_cap)

    doc.save(DOC_PATH)


if __name__ == "__main__":
    rebuild_section()
