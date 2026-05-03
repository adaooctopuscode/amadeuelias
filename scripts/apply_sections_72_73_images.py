#!/usr/bin/env python3
"""
Substitui imagens nas seções 7.2 (paredes), 7.2.2 (transversinas entre apoios),
7.2.3 (pilares) e 7.3 (infraestrutura), conforme prefixos da pasta do WhatsApp.

Infraestrutura: 55 arquivos na ordem do chat, 54 parágrafos com imagem no DOC atual —
aplica-se os primeiros 54 ficheiros (exclui-se o último da lista: 00000434).
"""

from __future__ import annotations

import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from lxml import etree

ROOT = "/Users/carlos_adao/Documents/2AEngenharia/Amadeu_Elias"
DOC_PATH = os.path.join(ROOT, "MC-AMADEU ELIAS R0.docx")
IMG_DIR = os.path.join(ROOT, "WhatsApp Chat - PONTE AMADEU")


def find_jpg(prefix: str) -> str:
    for name in sorted(os.listdir(IMG_DIR)):
        if name.startswith(prefix + "-") and name.lower().endswith((".jpg", ".jpeg", ".png")):
            return os.path.join(IMG_DIR, name)
    raise FileNotFoundError(prefix)


def replace_paragraph_image(paragraph: Paragraph, image_path: str, width_in: float = 5.5):
    p = paragraph._p
    for child in list(p):
        if etree.QName(child).localname != "pPr":
            p.remove(child)
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(width_in))
    pf = paragraph.paragraph_format
    if pf.alignment is None:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER


def img_paragraph_indices(doc: Document, start: int, end: int) -> list[int]:
    out = []
    for i in range(start, end + 1):
        if "blip" in doc.paragraphs[i]._p.xml:
            out.append(i)
    return out


def apply():
    doc = Document(DOC_PATH)

    # --- 7.2.1 Paredes (após deformações): índices fixos verificados no memorial ---
    paredes_img_idx = [1184, 1187, 1191, 1194]
    paredes_prefs = ["00000281", "00000282", "00000285", "00000286"]
    for idx, pr in zip(paredes_img_idx, paredes_prefs):
        replace_paragraph_image(doc.paragraphs[idx], find_jpg(pr))

    # --- 7.2.2 ---
    trans_prefs = [
        "00000290",
        "00000292",
        "00000293",
        "00000294",
        "00000295",
        "00000296",
        "00000298",
        "00000299",
        "00000300",
        "00000301",
        "00000302",
        "00000303",
        "00000305",
        "00000306",
        "00000307",
        "00000308",
        "00000309",
        "00000310",
        "00000312",
        "00000313",
        "00000314",
        "00000315",
        "00000316",
        "00000317",
    ]
    trans_idx = img_paragraph_indices(doc, 1197, 1275)
    assert len(trans_idx) == len(trans_prefs), (len(trans_idx), len(trans_prefs))
    for idx, pr in zip(trans_idx, trans_prefs):
        replace_paragraph_image(doc.paragraphs[idx], find_jpg(pr))

    # --- 7.2.3 ---
    pilar_prefs = (
        [f"{n:08d}" for n in range(320, 326)]
        + [f"{n:08d}" for n in range(327, 333)]
        + [f"{n:08d}" for n in range(334, 340)]
        + [f"{n:08d}" for n in range(341, 347)]
        + ["00000348", "00000349", "00000350", "00000352", "00000353", "00000354"]
        + [f"{n:08d}" for n in range(356, 362)]
    )
    pil_idx = img_paragraph_indices(doc, 1276, 1393)
    assert len(pil_idx) == len(pilar_prefs), (len(pil_idx), len(pilar_prefs))
    for idx, pr in zip(pil_idx, pilar_prefs):
        replace_paragraph_image(doc.paragraphs[idx], find_jpg(pr))

    # --- 7.3 (54 imagens; 55 ficheiros — último da lista não aplicado) ---
    infra_prefs = [
        "00000364",
        "00000365",
        "00000368",
        "00000369",
        "00000370",
        "00000371",
        "00000372",
        "00000373",
        "00000375",
        "00000376",
        "00000377",
        "00000378",
        "00000379",
        "00000381",
        "00000382",
        "00000384",
        "00000385",
        "00000386",
        "00000387",
        "00000388",
        "00000390",
        "00000391",
        "00000393",
        "00000395",
        "00000396",
        "00000397",
        "00000398",
        "00000399",
        "00000401",
        "00000402",
        "00000404",
        "00000405",
        "00000406",
        "00000407",
        "00000408",
        "00000410",
        "00000411",
        "00000413",
        "00000414",
        "00000415",
        "00000416",
        "00000417",
        "00000418",
        "00000420",
        "00000421",
        "00000423",
        "00000424",
        "00000425",
        "00000426",
        "00000427",
        "00000429",
        "00000430",
        "00000432",
        "00000433",
    ]
    infra_idx = img_paragraph_indices(doc, 1394, 1600)
    if len(infra_idx) != len(infra_prefs):
        print(
            f"AVISO: infra índices {len(infra_idx)} vs prefs {len(infra_prefs)}",
            file=sys.stderr,
        )
        raise SystemExit(1)
    for idx, pr in zip(infra_idx, infra_prefs):
        replace_paragraph_image(doc.paragraphs[idx], find_jpg(pr))

    # Legenda conforme WhatsApp: primeira figura do Encontro 1 inclui "Rz"
    old232 = "Figura 232: Reações nas estacas do Bloco do Encontro 1."
    new232 = "Figura 232: Reações Rz nas estacas do Bloco do Encontro 1."
    for para in doc.paragraphs:
        if para.text.strip() == old232:
            for child in list(para._p):
                if etree.QName(child).localname != "pPr":
                    para._p.remove(child)
            para.add_run(new232)
            break

    doc.save(DOC_PATH)
    print("OK: imagens 7.2 / 7.2.2 / 7.2.3 / 7.3 atualizadas.")


if __name__ == "__main__":
    apply()
