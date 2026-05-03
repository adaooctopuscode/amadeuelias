#!/usr/bin/env python3
"""
Secção 7.3 (INFRAESTRUTURA): prefixa itens de lista com a) … g) e corrige
legenda duplicada Figura 236 (segunda ocorrência → Figura 237).
"""
from __future__ import annotations

import os
import sys
from docx import Document
from docx.oxml.ns import qn

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOC_PATH = os.path.join(ROOT, "MC-AMADEU ELIAS R0.docx")

HEADER_MAP = [
    ("Bloco do Encontro 1:", "a) Bloco do Encontro 1:"),
    ("Blocos dos Pilares 1 e 2:", "b) Blocos dos Pilares 1 e 2:"),
    ("Blocos dos Pilares 3 e 4:", "c) Blocos dos Pilares 3 e 4:"),
    ("Blocos dos Pilares 5 e 6:", "d) Blocos dos Pilares 5 e 6:"),
    ("Blocos dos Pilares 7 e 8:", "e) Blocos dos Pilares 7 e 8:"),
    ("Blocos dos Pilares 9 e 10:", "f) Blocos dos Pilares 9 e 10:"),
    ("Bloco do Encontro 2:", "g) Bloco do Encontro 2:"),
]

DUP_CAPTION = "Figura 236: Reações nas estacas dos blocos dos Pilares 3 e 4."
FIX_CAPTION = "Figura 237: Reações nas estacas dos blocos dos Pilares 3 e 4."


def replace_paragraph_text_preserving_style(para, new_text: str) -> None:
    """Substitui o texto do parágrafo mantendo w:pPr (estilo, indentação de lista)."""
    p = para._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)
    para.add_run(new_text)


def main() -> None:
    doc = Document(DOC_PATH)
    n_headers = 0
    for old, new in HEADER_MAP:
        for para in doc.paragraphs:
            if para.text.strip() == old:
                replace_paragraph_text_preserving_style(para, new)
                n_headers += 1
                break
    if n_headers != len(HEADER_MAP):
        print(f"ERRO: cabeçalhos encontrados {n_headers}, esperado {len(HEADER_MAP)}", file=sys.stderr)
        raise SystemExit(1)

    seen_dup = 0
    for para in doc.paragraphs:
        if para.text.strip() != DUP_CAPTION:
            continue
        seen_dup += 1
        if seen_dup == 2:
            replace_paragraph_text_preserving_style(para, FIX_CAPTION)
            break
    else:
        print("ERRO: segunda legenda Figura 236 não encontrada", file=sys.stderr)
        raise SystemExit(1)

    doc.save(DOC_PATH)
    print("OK: 7.3 — itens a)–g) e legenda 237 aplicados.")
    print("Gravado:", DOC_PATH)


if __name__ == "__main__":
    main()
